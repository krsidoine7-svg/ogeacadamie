import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_runs_to_paragraph(p, text):
    tokens = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = p.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('`') and token.endswith('`'):
            run = p.add_run(token[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(199, 37, 78)
        else:
            p.add_run(token)

def add_paragraph_with_runs(doc, text, style=None, is_bullet=False):
    if is_bullet:
        p = doc.add_paragraph(style='List Bullet')
    else:
        p = doc.add_paragraph(style=style)
    add_runs_to_paragraph(p, text)
    return p

def create_roles_table(doc, roles_data):
    if not roles_data:
        return
        
    num_cols = 4
    num_rows = len(roles_data) + 1
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    headers = [
        "Rôle / Profil", 
        "Description & Périmètre", 
        "Droits & Fonctionnalités Autorisées", 
        "Actions Strictement Interdites"
    ]
    
    # Style header row
    hdr_row = table.rows[0]
    trPr = hdr_row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    
    for c_idx, text in enumerate(headers):
        cell = hdr_row.cells[c_idx]
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        set_cell_background(cell, "1A5276")
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    for r_idx, role in enumerate(roles_data):
        row = table.rows[r_idx + 1]
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        
        # Cell 0: Role Name
        cell_name = row.cells[0]
        set_cell_margins(cell_name, top=120, bottom=120, left=150, right=150)
        p = cell_name.paragraphs[0]
        run = p.add_run(role['name'])
        run.bold = True
        run.font.size = Pt(10.5)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Cell 1: Description
        cell_desc = row.cells[1]
        set_cell_margins(cell_desc, top=120, bottom=120, left=150, right=150)
        p = cell_desc.paragraphs[0]
        add_runs_to_paragraph(p, role['desc'])
        
        # Cell 2: Allowed Rights
        cell_allowed = row.cells[2]
        set_cell_margins(cell_allowed, top=120, bottom=120, left=150, right=150)
        p = cell_allowed.paragraphs[0]
        for idx, item in enumerate(role['allowed']):
            if idx > 0:
                p = cell_allowed.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
            add_runs_to_paragraph(p, "• " + item)
            
        # Cell 3: Forbidden Actions
        cell_forbidden = row.cells[3]
        set_cell_margins(cell_forbidden, top=120, bottom=120, left=150, right=150)
        p = cell_forbidden.paragraphs[0]
        for idx, item in enumerate(role['forbidden']):
            if idx > 0:
                p = cell_forbidden.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
            add_runs_to_paragraph(p, "• " + item)
            
        # Zebra striping
        if (r_idx + 1) % 2 == 0:
            for cell in row.cells:
                set_cell_background(cell, "F2F4F4")
                
    doc.add_paragraph().paragraph_format.space_before = Pt(6)

def create_word_table(doc, rows):
    if not rows:
        return
        
    parsed_rows = []
    for row in rows:
        cells = [c.strip() for c in row.split('|')]
        if row.startswith('|'):
            cells = cells[1:]
        if row.endswith('|'):
            cells = cells[:-1]
        parsed_rows.append(cells)
        
    num_cols = len(parsed_rows[0])
    num_rows = len(parsed_rows)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    
    for r_idx, row_data in enumerate(parsed_rows):
        row = table.rows[r_idx]
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        
        if r_idx == 0:
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
            
        for c_idx in range(min(num_cols, len(row_data))):
            cell = row.cells[c_idx]
            cell_text = row_data[c_idx]
            set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
            
            p = cell.paragraphs[0]
            add_runs_to_paragraph(p, cell_text)
            
            if r_idx == 0:
                set_cell_background(cell, "1A5276")
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                if r_idx % 2 == 0:
                    set_cell_background(cell, "F2F4F4")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
    doc.add_paragraph().paragraph_format.space_before = Pt(6)

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(33, 37, 41)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_rows = []
    
    in_rights_section = False
    roles_data = []
    current_role = None
    collect_mode = None
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Handle empty lines
        if not stripped:
            if in_table:
                create_word_table(doc, table_rows)
                table_rows = []
                in_table = False
            if in_rights_section:
                i += 1
                continue
            add_paragraph_with_runs(doc, "")
            i += 1
            continue
            
        # Rights section parsing mode
        if stripped.startswith('## 👥 2.'):
            in_rights_section = True
            h = doc.add_heading("2. DROITS ET FONCTIONNALITÉS PAR INTERVENANT", level=2)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.keep_with_next = True
            run = h.runs[0]
            run.font.name = 'Calibri Light'
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(26, 82, 118)
            i += 1
            continue
            
        if in_rights_section:
            if stripped.startswith('## 🧪 3.'):
                if current_role:
                    roles_data.append(current_role)
                    current_role = None
                create_roles_table(doc, roles_data)
                in_rights_section = False
                # Do not increment i, let normal parser catch the section 3 heading
            else:
                if stripped.startswith('### '):
                    if current_role:
                        roles_data.append(current_role)
                    role_title = stripped.lstrip('#').strip()
                    # Strip emojis and formatting
                    role_title = re.sub(r'[👤💼👑🔌]', '', role_title).strip()
                    current_role = {'name': role_title, 'desc': '', 'allowed': [], 'forbidden': []}
                    collect_mode = None
                elif current_role:
                    if stripped.startswith('**Description & Périmètre :**'):
                        collect_mode = 'desc'
                        desc_text = stripped[len('**Description & Périmètre :**'):].strip()
                        current_role['desc'] = desc_text
                    elif stripped.startswith('**Droits & Fonctionnalités Autorisées :**'):
                        collect_mode = 'allowed'
                    elif stripped.startswith('**Actions Strictement Interdites :**'):
                        collect_mode = 'forbidden'
                    elif stripped.startswith('- ') or stripped.startswith('* '):
                        bullet_item = stripped[2:].strip()
                        if collect_mode == 'allowed':
                            current_role['allowed'].append(bullet_item)
                        elif collect_mode == 'forbidden':
                            current_role['forbidden'].append(bullet_item)
                    else:
                        if collect_mode == 'desc':
                            current_role['desc'] += " " + stripped
                i += 1
                continue
                
        # Handle regular Table rows
        if stripped.startswith('|'):
            in_table = True
            if re.match(r'^\|[\s:-|]+$', stripped):
                i += 1
                continue
            table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            create_word_table(doc, table_rows)
            table_rows = []
            in_table = False
            
        # Handle Headings
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            title_text = stripped.lstrip('#').strip()
            clean_title = re.sub(r'\*\*|`', '', title_text)
            
            h = doc.add_heading(clean_title, level=min(level, 4))
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.keep_with_next = True
            
            run = h.runs[0]
            run.font.name = 'Calibri Light'
            if level == 1:
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor(10, 14, 23)
            elif level == 2:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(26, 82, 118)
            elif level == 3:
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(46, 64, 87)
            else:
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(50, 50, 50)
                
            i += 1
            continue
            
        # Handle List Bullets
        if stripped.startswith('* ') or stripped.startswith('- '):
            bullet_text = stripped[2:].strip()
            if bullet_text.startswith('> '):
                bullet_text = bullet_text[2:]
            add_paragraph_with_runs(doc, bullet_text, is_bullet=True)
            i += 1
            continue
            
        # Handle Blockquotes
        if stripped.startswith('> '):
            text = stripped[2:].strip()
            if text.startswith('[!'):
                alert_type = re.match(r'^\[!(.*?)\]', text)
                if alert_type:
                    type_str = alert_type.group(1)
                    text = text[len(alert_type.group(0)):].strip()
                    p = add_paragraph_with_runs(doc, f"[{type_str}] {text}")
                    p.paragraph_format.left_indent = Inches(0.5)
                    for run in p.runs:
                        run.italic = True
                        run.font.color.rgb = RGBColor(120, 120, 120)
            else:
                p = add_paragraph_with_runs(doc, text)
                p.paragraph_format.left_indent = Inches(0.5)
                for run in p.runs:
                    run.italic = True
                    run.font.color.rgb = RGBColor(100, 100, 100)
            i += 1
            continue
            
        # Normal paragraph
        add_paragraph_with_runs(doc, stripped)
        i += 1

    if in_table and table_rows:
        create_word_table(doc, table_rows)
        
    doc.save(docx_path)
    print(f"File successfully converted to {docx_path}")

if __name__ == '__main__':
    md_file = "c:\\Users\\Toto.ADMINISTRATOR\\Desktop\\oge-academie\\cahier_de_recette.md"
    docx_file = "c:\\Users\\Toto.ADMINISTRATOR\\Desktop\\oge-academie\\cahier_de_recette.docx"
    convert_md_to_docx(md_file, docx_file)
